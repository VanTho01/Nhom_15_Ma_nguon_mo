import sympy as sym
from customtkinter import *
from tkinter import messagebox, filedialog
from PIL import Image
import os
from docx import Document
import math
import subprocess

class calLimit:
    def __init__(self, win):
        

        self.win = win
        win.configure(fg_color = '#B48989')
        win.geometry('1560x988')

        CTkLabel(win, corner_radius=10, text = "",fg_color = '#9FE1FD', width = 1515, height = 858).place(x = 23, y = 113)
        CTkLabel(win, corner_radius=10, text = "", width = 1042, fg_color = '#CEF7E6', bg_color = '#9FE1FD', height = 774).place(x = 277, y = 150)
        CTkLabel(win, text = "", fg_color = "#A56666", width = 96, height = 90).place(x = 1300, y = 18)


        CTkLabel(win, corner_radius=8, text = "Chọn Cận", font = ('Times new roman', 25), fg_color = '#C59FF5', bg_color = '#CEF7E6', text_color = 'Black', width = 263, height = 70).place(x = 374, y = 305)

        self.selected_option = StringVar()
        self.selected_option.set('-')
        self.options = ['-','+']
        self.option_menu = CTkOptionMenu(win, font = ('Times new roman', 30), bg_color = '#CEF7E6', corner_radius = 10, variable = self.selected_option,  values = self.options, hover = True, button_hover_color = '#4287f5', text_color = 'black', width = 172, height = 70)
        self.option_menu.place(x = 1038, y = 305)

        self.radio_var1 = StringVar()
        self.radio_var1.set('cận 1 phía')
        self.options2 = ['cận 1 phía', 'cận 2 phía']
        self.can_value = CTkOptionMenu(win, font = ('Times new roman', 25), command = self.selection, bg_color = '#CEF7E6', corner_radius= 10, variable = self.radio_var1,  values = self.options2, hover = True, button_hover_color = '#4287f5', text_color = 'black', width = 276, height = 70)
        self.can_value.place(x = 658, y = 305)

        self.expression_label = CTkLabel(win, corner_radius = 10, text = "Hàm f(x)", font = ('Times new roman', 25), fg_color = '#C59FF5', bg_color = '#CEF7E6', text_color = 'Black', width = 263, height = 70)
        self.expression_label.place(x = 374, y = 192)
        self.expression_entry = CTkTextbox(win, bg_color = '#CEF7E6', corner_radius=6, font = ('Times new roman', 25), text_color = 'Black', width = 552, height = 70)
        self.expression_entry.place(x = 658, y = 192)

        self.output = CTkTextbox(win, bg_color = '#CEF7E6', corner_radius = 10, font = ('Times New Roman', 25), height = 309, width = 552)
        self.output.configure(text_color = 'black', fg_color = "white")
        self.output.place(x = 658, y = 531)

        self.file = CTkTextbox(win, bg_color = '#CEF7E6', corner_radius = 10, font = ('Times New Roman', 25), height = 82, width = 186)
        self.file.configure(text_color = 'black', fg_color = "white", bg_color = '#9FE1FD')
        self.file.place(x = 1338, y = 152)

        #expression = expression_entry.get()
        self.desti_label = CTkLabel(win, corner_radius= 10, text = "Giá trị limit", font = ('Times new roman', 25), fg_color = '#C59FF5', bg_color = '#CEF7E6', text_color = 'Black', width = 263, height = 70)
        self.desti_label.place(x = 374, y = 418)
        self.desti_value = CTkTextbox(win, bg_color = '#CEF7E6', corner_radius= 10, font = ('Times new roman', 25), text_color = 'Black', fg_color = 'white', width = 276, height = 70)
        self.desti_value.place(x = 658, y = 418)
        #type_can = on_radio_button_select()

        self.limit_button = CTkButton(win, corner_radius = 10, text = "Kết Quả", command = self.compute_Limit, font = ('Times new roman', 25), fg_color = '#C59FF5', bg_color = '#CEF7E6', 
                                      text_color = 'Black', width = 263, height = 110, hover = True, hover_color='#37a4de', border_color = '#37a4de', border_width = 3)
        self.limit_button.place(x = 374, y = 531)

        self.exit_button = CTkButton(win, corner_radius = 10, text = "Save", command = self.storeResult, font = ('Times new roman', 25), fg_color = '#C6E8B1', bg_color = '#9FE1FD', text_color = 'Black', 
                                     width = 217, height = 89, hover = True, hover_color='#37a4de', border_color = '#37a4de', border_width = 3)
        self.exit_button.place(x = 39, y = 261)

        self.iconUpload = CTkImage(light_image = Image.open('asset/iconUpload.jpg'), size = (83, 77))
        self.upload = CTkButton(win, image = self.iconUpload, corner_radius = 3, text = "", command = self.upFile, font = ('Times new roman', 25), fg_color = "#54AFC3", 
                                text_color = 'Black', width = 82, height = 76, hover = True, hover_color='#37a4de', border_color = '#37a4de', border_width = 3)
        self.upload.place(x = 1422, y = 18)

        self.help_button = CTkButton(win, corner_radius = 10, text = "Help", font = ('Times new roman', 25), fg_color = '#C6E8B1', bg_color = '#9FE1FD', text_color = 'Black', width = 217, 
                                    command=self.help, height = 89, hover = True, hover_color='#37a4de', border_color = '#37a4de', border_width = 3)
        self.help_button.place(x = 39, y = 148)


        if os.path.isfile("calLimit/expression.txt") == True and os.path.isfile("calLimit/resultLimit.txt") == True and os.path.isfile("calLimit/destiValue.txt") == True:
            pre_express = open("calLimit/expression.txt", "r")
            self.expression_entry.insert(INSERT, pre_express.read())
            pre_result = open("calLimit/resultLimit.txt", "r")
            text="Giới hạn của hàm số: " + pre_result.read()
            self.output.insert(INSERT, text)
            desti = open('calLimit/destiValue.txt', 'r')
            self.desti_value.insert(INSERT, desti.read())
        if os.path.isfile('calLimit/fileStore.txt') == True:
            file1 = open('calLimit/fileStore.txt', 'r')
            self.file.insert(INSERT, file1.read())
    def selection(self, event):
        self.option_menu.configure(state = "disabled" if self.radio_var1.get() == 'cận 2 phía' else "normal")
    def compute_Limit(self):
        try:
            expr = open("calLimit/expression.txt", "w")
            result = open("calLimit/resultLimit.txt", "w")
            desti = open("calLimit/destiValue.txt", 'w')
            self.output.delete("1.0", "end")
            type_can = self.radio_var1.get()
            x = sym.symbols('x')
            expression = self.expression_entry.get("1.0", "end")
            expr.write(expression)
            can = self.desti_value.get("1.0", "end")
            if (str(type_can) == 'cận 1 phía'):
                if can != "-oo\n" and can != "oo\n":
                    integral = sym.limit(expression, x, int(can), str(self.selected_option.get()))
                else: 
                    messagebox.showwarning('Lỗi chọn giá trị tiến tới!', 'không  tồn tại giá trị vô cực giới hạn 1 phía!')
            if (str(type_can) == 'cận 2 phía'):
                if can != "-oo\n" and can != "oo\n" and can != "+oo\n":
                    integral = sym.limit(expression, x, int(can))
                else:
                    if can == "-oo\n":
                        integral = sym.limit(expression, x, -math.inf)
                    if can == "+oo\n" or can == "oo\n":
                        integral = sym.limit(expression, x, +math.inf)
            text="Giới hạn của hàm số: " + str(integral)
            self.output.insert(INSERT, text)
            result.write(str(integral))
            desti.write(str(can))
        except:
            messagebox.showerror('Error', 'Kiểm tra lại đầu vào!')
    def upFile(self):
        file_path = filedialog.askopenfilename()
        if file_path:
            file_name = file_path.split("/")[-1]
            if (file_name.split(".")[-1] != "txt"):
                messagebox.showwarning('Định dạng file không đúng', 'File được upload phải có định dạng .txt')
            else:
                express = open(file_path, "r")
                fileSTr = open('calLimit/fileStore.txt', 'w')
                express = open(file_path, "r")

                self.file.delete("1.0", "end")
                self.file.insert(INSERT, file_name)

                self.expression_entry.delete("1.0", "end")
                self.expression_entry.insert(INSERT, express.read().strip())
                fileSTr.write(file_name)

    def storeResult(self):
        try:
            doc = Document()
            text1 = f'-- Hàm số: {self.expression_entry.get("1.0", "end")}'
            doc.add_paragraph(text1)
            text2 = f'-- Giá trị tiến tới: {self.desti_value.get("1.0", "end")}'
            doc.add_paragraph(text2)
            text3 = f'-- {self.output.get("1.0", "end")}'
            doc.add_paragraph(text3)
            file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")], title="Save As")
            try:
                doc.save(f'{file_path}')
                messagebox.showinfo('Lưu file thành công!.',f'- Địa chỉ file: {file_path}')
            except PermissionError:
                messagebox.showinfo('Lưu file không thành công!', 'File đang được mở! Hãy tắt file đi!')
        except:
            messagebox.showwarning('Lưu file thất bại!', 'Xem lại dữ liệu đầu vào!')
    def help(self):
        file = "Help/HelpLimit/helpLimit.py"
        subprocess.call(['python', file])

if __name__ == "__main__":
    win = CTk()
    app = calLimit(win)
    win.mainloop()
