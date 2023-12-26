import sympy as sym
import sympy.plotting as syp
from customtkinter import *
from tkinter import messagebox, filedialog
from PIL import Image
import os
from docx import Document
import math
import subprocess
from docx.shared import Inches

class calLimit:
    def __init__(self, win):
        

        self.win = win
        win.configure(fg_color = '#B48989')
        win.geometry('1560x988')

        CTkLabel(win, corner_radius=10, text = "",fg_color = '#9FE1FD', width = 1515, height = 858).place(x = 23, y = 113)
        CTkLabel(win, corner_radius=10, text = "", width = 695, fg_color = '#CEF7E6', bg_color = '#9FE1FD', height = 686).place(x = 46, y = 135)
        CTkLabel(win, text = "", fg_color = "#A56666", width = 96, height = 90).place(x = 1300, y = 18)


        CTkLabel(win, corner_radius=8, text = "Chọn Cận", font = ('Times new roman', 25), fg_color = '#C59FF5', bg_color = '#CEF7E6', text_color = 'Black', width = 174, height = 84).place(x = 72, y = 272)

        self.selected_option = StringVar()
        self.selected_option.set('-')
        self.options = ['-','+']
        self.option_menu = CTkOptionMenu(win, font = ('Times new roman', 30), bg_color = '#CEF7E6', corner_radius = 10, variable = self.selected_option,  values = self.options, hover = True, button_hover_color = '#4287f5', text_color = 'black', width = 172, height = 83)
        self.option_menu.place(x = 532, y = 272)

        self.radio_var1 = StringVar()
        self.radio_var1.set('cận 1 phía')
        self.options2 = ['cận 1 phía', 'cận 2 phía']
        self.can_value = CTkOptionMenu(win, font = ('Times new roman', 25), command = self.selection, bg_color = '#CEF7E6', corner_radius= 10, variable = self.radio_var1,  values = self.options2, hover = True, button_hover_color = '#4287f5', text_color = 'black', width = 242, height = 84)
        self.can_value.place(x = 267, y = 272)

        self.expression_label = CTkLabel(win, corner_radius = 10, text = "Hàm f(x)", font = ('Times new roman', 25), fg_color = '#C59FF5', bg_color = '#CEF7E6', text_color = 'Black', width = 174, height = 84)
        self.expression_label.place(x = 72, y = 155)
        self.expression_entry = CTkTextbox(win, bg_color = '#CEF7E6', corner_radius=6, font = ('Times new roman', 25), text_color = 'Black', width = 437, height = 84)
        self.expression_entry.place(x = 267, y = 155)

        self.output = CTkTextbox(win, bg_color = '#CEF7E6', corner_radius = 10, font = ('Times New Roman', 25), height = 208, width = 437)
        self.output.configure(text_color = 'black', fg_color = "white")
        self.output.place(x = 267, y = 494)

        self.file = CTkTextbox(win, bg_color = '#CEF7E6', corner_radius = 10, font = ('Times New Roman', 25), height = 82, width = 632)
        self.file.configure(text_color = 'black', fg_color = "white", bg_color = '#CEF7E6')
        self.file.place(x = 72, y = 720)

        #expression = expression_entry.get()
        self.desti_label = CTkLabel(win, corner_radius= 10, text = "Giá trị limit", font = ('Times new roman', 25), fg_color = '#C59FF5', bg_color = '#CEF7E6', text_color = 'Black', width = 174, height = 84)
        self.desti_label.place(x = 72, y = 382)
        self.desti_value = CTkTextbox(win, bg_color = '#CEF7E6', corner_radius= 10, font = ('Times new roman', 25), text_color = 'Black', fg_color = 'white', width = 242, height = 84)
        self.desti_value.place(x = 267, y = 382)
        #type_can = on_radio_button_select()

        self.limit_button = CTkButton(win, corner_radius = 10, text = "Kết Quả", command = self.compute_Limit, font = ('Times new roman', 25), fg_color = '#C59FF5', bg_color = '#CEF7E6', 
                                      text_color = 'Black', width = 174, height = 96, hover = True, hover_color='#37a4de', border_color = '#37a4de', border_width = 3)
        self.limit_button.place(x = 72, y = 494)

        self.graph_button = CTkButton(win, corner_radius = 10, text = "Đồ thị", command = self.draw_graph, font = ('Times new roman', 25), fg_color = '#C59FF5', bg_color = '#9FE1FD', 
                                      text_color = 'Black', width = 174, height = 96, hover = True, hover_color='#37a4de', border_color = '#37a4de', border_width = 3)
        self.graph_button.place(x = 1050, y = 856)

        self.exit_button = CTkButton(win, corner_radius = 10, text = "Lưu kết quả", command = self.storeResult, font = ('Times new roman', 25), fg_color = '#C6E8B1', bg_color = '#9FE1FD', text_color = 'Black', 
                                     width = 217, height = 89, hover = True, hover_color='#37a4de', border_color = '#37a4de', border_width = 3)
        self.exit_button.place(x = 524, y = 863)

        self.iconUpload = CTkImage(light_image = Image.open('asset/iconUpload.jpg'), size = (83, 77))
        self.upload = CTkButton(win, image = self.iconUpload, corner_radius = 3, text = "", command = self.upFile, font = ('Times new roman', 25), fg_color = "#54AFC3", 
                                text_color = 'Black', width = 82, height = 76, hover = True, hover_color='#37a4de', border_color = '#37a4de', border_width = 3)
        self.upload.place(x = 1422, y = 18)

        self.help_button = CTkButton(win, corner_radius = 10, text = "Hướng dẫn", font = ('Times new roman', 25), fg_color = '#C6E8B1', bg_color = '#9FE1FD', text_color = 'Black', width = 217, 
                                    command=self.help, height = 89, hover = True, hover_color='#37a4de', border_color = '#37a4de', border_width = 3)
        self.help_button.place(x = 46, y = 863)

        self.graph_area = CTkLabel(win, corner_radius=10, text = "", width = 765, fg_color = '#CEF7E6', bg_color = '#9FE1FD', height = 686)
        self.graph_area.place(x = 754, y = 135)
        if os.path.isfile('calLimit/dothi.png') == True:
            picture_graph = CTkImage(light_image = Image.open('calLimit/dothi.png'), size = (765, 686))
            self.graph_area.configure(image = picture_graph)
        if os.path.isfile("calLimit/expression.txt") == True and os.path.isfile("calLimit/resultLimit.txt") == True and os.path.isfile("calLimit/destiValue.txt") == True:
            pre_express = open("calLimit/expression.txt", "r")
            self.expression_entry.insert(INSERT, pre_express.read())
            pre_result = open("calLimit/resultLimit.txt", "r")
            text="Giới hạn của hàm số: " + pre_result.read()
            self.output.insert(INSERT, text)
            desti = open('calLimit/destiValue.txt', 'r')
            self.desti_value.insert(INSERT, desti.read())
        if os.path.isfile('calLimit/fileStore.txt') == True:
            file1 = open('calLimit/fileStore.txt', 'r')
            self.file.insert(INSERT, file1.read())
    def selection(self, event):
        self.option_menu.configure(state = "disabled" if self.radio_var1.get() == 'cận 2 phía' else "normal")
    def compute_Limit(self):
        try:
            expr = open("calLimit/expression.txt", "w")
            result = open("calLimit/resultLimit.txt", "w")
            desti = open("calLimit/destiValue.txt", 'w')
            self.output.delete("1.0", "end")
            type_can = self.radio_var1.get()
            x = sym.symbols('x')
            expression = self.expression_entry.get("1.0", "end")
            expr.write(expression)
            can = self.desti_value.get("1.0", "end")
            if (str(type_can) == 'cận 1 phía'):
                if can != "-oo\n" and can != "oo\n" and can != "+oo":
                    can = str(can).replace("\n", "")
                    can = str(can).replace("\n", "")
                    can = str(can).replace("pi", "math.pi")
                    can = str(can).replace("e", "math.e")
                    integral = sym.limit(expression, x, float(eval(can)))
                    integral = sym.limit(expression, x, float(eval(can)), str(self.selected_option.get()))
                else: 
                    messagebox.showwarning('Lỗi chọn giá trị tiến tới!', 'không  tồn tại giá trị vô cực giới hạn 1 phía!')
            if (str(type_can) == 'cận 2 phía'):
                if can != "-oo\n" and can != "oo\n" and can != "+oo\n":
                    can = str(can).replace("\n", "")
                    can = str(can).replace("pi", "math.pi")
                    can = str(can).replace("e", "math.e")
                    integral = sym.limit(expression, x, float(eval(can)))
                else:
                    if can == "-oo\n":
                        integral = sym.limit(expression, x, -math.inf)
                    if can == "+oo\n" or can == "oo\n":
                        integral = sym.limit(expression, x, +math.inf)
            text="Giới hạn của hàm số: " + str(integral)
            self.output.insert(INSERT, text)
            result.write(str(integral))
            desti.write(str(can))
        except:
            messagebox.showerror('Error', 'Kiểm tra lại đầu vào!')
    def upFile(self):
        file_path = filedialog.askopenfilename()
        if file_path:
            file_name = file_path.split("/")[-1]
            if (file_name.split(".")[-1] != "txt"):
                messagebox.showwarning('Định dạng file không đúng', 'File được upload phải có định dạng .txt')
            else:
                express = open(file_path, "r")
                fileSTr = open('calLimit/fileStore.txt', 'w')
                express = open(file_path, "r")

                self.file.delete("1.0", "end")
                self.file.insert(INSERT, file_name)

                self.expression_entry.delete("1.0", "end")
                self.expression_entry.insert(INSERT, express.read().strip())
                fileSTr.write(file_name)

    def storeResult(self):
        try:
            doc = Document()
            text1 = f'-- Hàm số: {self.expression_entry.get("1.0", "end")}'
            doc.add_paragraph(text1)
            text2 = f'-- Giá trị tiến tới: {self.desti_value.get("1.0", "end")}'
            doc.add_paragraph(text2)
            text3 = f'-- {self.output.get("1.0", "end")}'
            doc.add_paragraph(text3)
            text4 = "-- Đồ thị hàm số: "
            title = doc.add_paragraph(text4)
            image_path = 'calLimit/dothi.png'   
            run = title.add_run()
            run.add_picture(image_path, width=Inches(6), height=Inches(5))
            file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")], title="Save As")
            try:
                doc.save(f'{file_path}')
                messagebox.showinfo('Lưu file thành công!.',f'- Địa chỉ file: {file_path}')
            except PermissionError:
                messagebox.showinfo('Lưu file không thành công!', 'File đang được mở! Hãy tắt file đi!')
        except:
            messagebox.showwarning('Lưu file thất bại!', 'Xem lại dữ liệu đầu vào!')
    def draw_graph(self):
        try:
            result = open("calLimit/resultLimit.txt", "r").read().strip()
            x = sym.symbols('x')
            graph = syp.plot(sym.sympify(self.expression_entry.get("1.0", "end")), (x, -20, 20), title='Đồ thị hàm số f(x)', show = False)
            if result != '-oo' and result != 'oo' and result != '+oo':
                value_limit = syp.plot(sym.sympify(result),  (x, -20, 20), line_color='g',label = 'giá trị giới hạn', show = False)
                graph.append(value_limit[0])
            graph.legend = True
            graph.show()
            graph.save('calLimit/dothi.png')
            picture_graph = CTkImage(light_image = Image.open('calLimit/dothi.png'), size = (765, 686))
            self.graph_area.configure(image = picture_graph)
        except:
            if TypeError or SyntaxError:
                messagebox.showerror('Error', 'Kiểm tra lại đầu vào!')
            if ValueError:
                messagebox.showwarning('Warning!', 'Phương trình chỉ nhận một biến đầu vào là x!')
    def help(self):
        file = "Help/HelpLimit/helpLimit.py"
        subprocess.call(['python', file])

if __name__ == "__main__":
    win = CTk()
    app = calLimit(win)
    win.mainloop()
