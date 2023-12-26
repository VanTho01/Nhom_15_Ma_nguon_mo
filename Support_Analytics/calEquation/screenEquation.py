import sympy as sym
from sympy import *
from sympy.solvers.diophantine import diop_solve
from customtkinter import *
from tkinter import messagebox
import sympy.plotting as syp
from PIL import Image
import os
from docx import Document
from docx.shared import Inches
import subprocess


class calEquation:
    def __init__(self, win):
        self.win = win
        win.title('Muc 3')
        win.geometry('1560x988')
        win.configure(fg_color = '#B48989')

        CTkLabel(win, text = "", fg_color = '#76BFF4', width = 1515, height = 858, corner_radius=10).place(x = 23, y = 113)
        CTkLabel(win, text = "", fg_color = "#A56666", width = 96, height = 90).place(x = 1300, y = 18)

        self.control_area = CTkLabel(win,  corner_radius=10, fg_color = '#D9D9D9', bg_color = '#76BFF4', text = "", width = 592, height = 679)
        self.control_area.place(x = 60, y = 131)

        self.graph_area = CTkLabel(win,  corner_radius=10, fg_color = '#D9D9D9', bg_color = '#76BFF4', text = "", width = 851, height = 814)
        self.graph_area.place(x = 664, y = 131)

        if os.path.isfile('calEquation/dothi.png') == True:
            picture_graph = CTkImage(light_image = Image.open('calEquation/dothi.png'), size = (851, 814))
            self.configure(image = picture_graph)

        self.typePtr = StringVar()
        self.typePtr.set("phương trình 1 biến")
        options2 = ["phương trình 1 biến", "phương trình 2 biến diophantine"]
        self.can_value = CTkOptionMenu(win, font = ('Times new roman', 25), corner_radius= 10, variable = self.typePtr,  values = options2, 
                                    bg_color = "#D9D9D9", fg_color = '#F8F3F9', hover = True, button_hover_color = '#4287f5', text_color = 'black', width = 518, height = 68)
        self.can_value.place(x = 93, y = 153)

        self.upload_area = CTkTextbox(win, font = ('Times New Roman', 25), height = 68, width = 526, text_color = 'black', corner_radius=10)
        self.upload_area.configure(fg_color = '#F8F3F9', bg_color = '#D9D9D9')
        self.upload_area.place(x = 93, y = 700)

        self.iconUpload = CTkImage(light_image = Image.open('asset/iconUpload.jpg'), size = (83, 77))
        self.upload = CTkButton(win, image = self.iconUpload, corner_radius = 3, text = "", command = self.upFile, font = ('Times new roman', 25), fg_color = "#54AFC3", 
                                text_color = 'Black', width = 82, height = 76, hover = True, hover_color='#37a4de', border_color = '#37a4de', border_width = 3)
        self.upload.place(x = 1422, y = 18)

        self.expression_label = CTkLabel(win, text = "F(x)", font = ('Times new roman', 25), fg_color = '#E0F1C9', bg_color = '#D9D9D9', text_color = 'black', width = 114, height = 68, corner_radius=10)
        self.expression_label.place(x = 93, y = 261)
        self.expression_entry = CTkTextbox(win, font = ('Times new roman', 25), fg_color = '#F8F3F9', bg_color = '#D9D9D9', text_color = 'black', width = 392, height = 68, corner_radius=10)
        self.expression_entry.place(x = 227, y = 261)

        self.sovle_button = CTkButton(win, border_color = '#12151c', border_width = 1, text = " Giải \nptr",command = self.sovle_equation, font = ('Times new roman', 25), 
                                      fg_color = '#E0F1C9', bg_color = '#D9D9D9', text_color = 'black', corner_radius = 10, 
                    hover_color = '#fcc603', width = 114, height = 93)
        self.sovle_button.place(x = 93, y = 372)

        self.draw_button = CTkButton(win, border_color = '#12151c', border_width = 1, text = " Vẽ \nđồ thị ",command = self.draw_graph, font = ('Times new roman', 25), 
                                     fg_color = '#E0F1C9', bg_color = '#D9D9D9', text_color = 'black', corner_radius = 10, 
                    hover_color = '#fcc603', width = 114, height = 93)
        self.draw_button.place(x = 93, y = 564)

        self.save_button = CTkButton(win, border_color = '#12151c', border_width = 1, text = " Lưu kết quả ", command = self.storeResult, font = ('Times new roman', 25), fg_color = '#C6E8B1',  
                                     text_color = 'black', corner_radius = 10, bg_color = '#76BFF4',
                    hover_color = '#6f8fed', width = 217, height = 89)
        self.save_button.place(x = 435, y = 856)

        self.help_button = CTkButton(win, border_color = '#12151c', border_width = 1, text = " Hướng dẫn ",command = self.help, font = ('Times new roman', 25), fg_color = '#C6E8B1',  
                                     text_color = 'black', corner_radius = 10, bg_color = '#76BFF4',
                    hover_color = '#6f8fed', width = 217, height = 89)
        self.help_button.place(x = 64, y = 856)

        self.output = CTkTextbox(win, font = ('Times New Roman', 25), height = 285, width = 392, text_color = 'black', corner_radius=10)
        self.output.configure(fg_color = '#F8F3F9', bg_color = '#D9D9D9')
        self.output.place(x = 227, y = 372)
        if  os.path.isfile('calEquation/expression.txt') == True: 
            expr = open('calEquation/expression.txt', 'r')
            self.expression_entry.insert(INSERT, expr.read()) 
        if os.path.isfile('calEquation/resultEquation.txt') == True:
            reslt = open('calEquation/resultEquation.txt', 'r')
            text1 = "\n  Nghiệm của phương trình:\n " + reslt.read()
            self.output.insert(INSERT, text1)
        if os.path.isfile('calEquation/fileStore.txt') == True:
            file = open('calEquation/fileStore.txt', 'r')
            self.upload_area.insert(INSERT, file.read())

    def upFile(self):
        file_path = filedialog.askopenfilename()
        if file_path:
            file_name = file_path.split("/")[-1]
            if (file_name.split(".")[-1] != "txt"):
                messagebox.showwarning('Định dạng file không đúng', 'File được upload phải có định dạng .txt')
            else:
                try:
                    express = open(file_path, "r").read()
                    fileSTr = open('calEquation/fileStore.txt', 'w')
                    self.upload_area.delete("1.0", "end")
                    self.upload_area.insert(INSERT, file_name)
                    if express.strip() == "":
                        messagebox.showwarning('File có vấn đề!', 'File rỗng!')
                    self.expression_entry.delete("1.0", "end")
                    self.expression_entry.insert(INSERT, express.strip())
                    fileSTr.write(str(file_name))
                except:
                    if UnicodeDecodeError:
                        messagebox.showwarning("Lỗi File!", "Không đọc được nội dung trong file!")
            
    def sovle_equation(self):
        try:
            expr = open('calEquation/expression.txt', 'w')
            reslt = open('calEquation/resultEquation.txt', 'w')
            self.output.delete("1.0", "end")
            typePtr = self.typePtr.get()
            equation = self.expression_entry.get("1.0", "end")
            if str(typePtr) ==  'phương trình 1 biến':
                x = sym.symbols('x')
                self.solutions = sym.solve(sym.Eq(sym.sympify(equation), 0), x)
                text="\n  Nghiệm của phương trình:\n " + str(self.solutions)
            if str(typePtr) ==  'phương trình 2 biến diophantine':
                x, y = sym.symbols('x y')
                self.solutions = diophantine(sym.sympify(equation))
                if str(self.solutions) == 'set()' or str(self.solutions) == '(None, None)':
                    self.solutions = str("Root Empty")
                    text=f"\n  Nghiệm của phương trình:\n {str(self.solutions)}"
                else:
                    text=f"\n  Nghiệm của phương trình:\n {str(self.solutions)} \n với t thuộc Z(nếu có)"
            self.output.insert(INSERT, text)
            expr.write(str(equation)) 
            reslt.write(str(self.solutions))
        except:
            if SyntaxError or TypeError or sym.core.sympify.SympifyError:
                messagebox.showerror('Error', 'Kiểm tra lại đầu vào!')
            if RecursionError:
                messagebox.showwarning('Giải phương trình thất bại!', 'Không giải được dạng phương trình này!')
    def draw_graph(self):
        typePtr = self.typePtr.get()
        if str(typePtr) ==  'phương trình 1 biến':
            try:
                x = sym.symbols('x')
                self.graph = syp.plot(self.expression_entry.get("1.0", "end"), (x, -20, 20), title='Đồ thị hàm số f(x)').save('calEquation/dothi.png')
                picture_graph = CTkImage(light_image = Image.open('calEquation/dothi.png'), size = (851, 814))
                self.graph_area = CTkLabel(self.win,  corner_radius=10, fg_color = '#D9D9D9', text = "", width = 851, height = 814, image = picture_graph)
                self.graph_area.place(x = 664, y = 131)
            except:
                if TypeError or SyntaxError:
                    messagebox.showerror('Error', 'Kiểm tra lại đầu vào!')
                if ValueError:
                    messagebox.showwarning('Warning', 'Phương trình chỉ nhận một biến đầu vào là x!')
        if str(typePtr) ==  'phương trình 2 biến diophantine':
            try:
                x, y = sym.symbols('x y')
                self.graph = syp.plot3d(self.expression_entry.get("1.0", "end"), (x, -1.5, 1.5), (y, -1.5,1.5), title='Đồ thị hàm số f(x)').save('calEquation/dothi.png')
                picture_graph = CTkImage(light_image = Image.open('calEquation/dothi.png'), size = (851, 814))
                self.graph_area.configure(image = picture_graph)
            except:
                if TypeError or SyntaxError:
                    messagebox.showerror('Error', 'Kiểm tra lại đầu vào!')
                if ValueError:
                    messagebox.showwarning('Warning', 'Phương trình chỉ nhận 2 biến đầu vào là x và y!')

    def storeResult(self):
        try:
            doc = Document()
            text1 = f'-- Phuong trinh {self.expression_entry.get("1.0", "end")}'
            doc.add_paragraph(str(text1))
            doc.add_paragraph('\n')
            text2 = f'-- {self.output.get("1.0", "end")}'
            doc.add_paragraph(str(text2))
            doc.add_paragraph('\n')
            text3 = "-- Do thi ham so: "
            title = doc.add_paragraph(text3)
            image_path = 'calEquation/dothi.png'   
            run = title.add_run()
            run.add_picture(image_path, width=Inches(6), height=Inches(5))
            file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")], title="Save As")
            try:
                doc.save(f'{file_path}')
                messagebox.showinfo('Lưu file thành công!.',f'- Địa chỉ file: {file_path}')
            except PermissionError:
                messagebox.showinfo('Lưu file không thành công!', 'File đang được mở! Hãy tắt file đi!')
        except FileNotFoundError:
            messagebox.showwarning('Lưu file thất bại!', 'File kết quả không tồn tại!')
    def help(self):
        file = "Help/HelpEquation/helpEquation.py"
        subprocess.call(['python', file])

if __name__ == "__main__":
    win = CTk()
    app = calEquation(win)
    win.mainloop()