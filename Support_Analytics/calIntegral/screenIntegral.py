import sympy as sym
from customtkinter import *
from tkinter import messagebox, filedialog
import os
from PIL import Image
from docx import Document
from sympy import *
import subprocess
import math

class calIntegral:
    def __init__(self, win):
        self.win = win
        win.configure(fg_color = '#B48989')
        win.geometry('1560x988')

        CTkLabel(win, corner_radius=10, text = "",fg_color = '#B5EBF7', width = 1515, height = 858).place(x = 23, y = 113)
        CTkLabel(win, corner_radius=10, text = "", fg_color = '#CBD3F9', bg_color = '#B5EBF7', width = 1042, height = 774).place(x = 277, y = 150)
        CTkLabel(win, text = "", fg_color = "#A56666", width = 96, height = 90).place(x = 1300, y = 18)
        
        self.expression_label = CTkLabel(win, bg_color = '#CBD3F9', text = "f(x)", font = ('Times new roman', 25), fg_color = "#D465F0", text_color = "black", width = 165, height = 68, corner_radius = 10)
        self.expression_label.place(x = 341, y = 207)
        self.expression_entry = CTkTextbox(win, bg_color = '#CBD3F9', font = ('Times new roman', 23), fg_color = '#F8F3F9', text_color = 'black', width = 552, height = 68, corner_radius = 10)
        self.expression_entry.place(x = 554, y = 207)

        self.integral_type_label = CTkLabel(win, bg_color = '#CBD3F9', text = "Loại", font = ('Times new roman', 25), corner_radius = 10, fg_color = "#D465F0", text_color = "black", 
                                            width = 165, height = 68)
        self.integral_type_label.place(x=341, y = 296)

        self.radio_var1 = StringVar()
        self.radio_var1.set("type?")
        options2 = ["definite", "indefinite"]
        self.can_value = CTkOptionMenu(win, bg_color = '#CBD3F9', font = ('Times new roman', 25), command = self.selection, corner_radius= 10, variable = self.radio_var1,  values = options2, 
                                       fg_color = '#F8F3F9', hover = True, button_hover_color = '#4287f5', text_color = 'black', width = 276, height = 70)
        self.can_value.place(x = 554, y = 296)

        self.level_label = CTkLabel(win, bg_color = '#CBD3F9', text = "Cấp", font = ('Times new roman', 25), corner_radius = 10, fg_color = "#D465F0", text_color = "black", 
                                            width = 118, height = 68)
        self.level_label.place(x=410, y = 476)
        self.radio_var2 = StringVar()
        self.radio_var2.set("cấp 1")
        options3 = ["cấp 1", "cấp 2", "cấp 3"]
        self.cap_value = CTkOptionMenu(win, bg_color = '#CBD3F9', font = ('Times new roman', 25), corner_radius= 10, variable = self.radio_var2,  values = options3, 
                                       fg_color = '#F8F3F9', hover = True, button_hover_color = '#4287f5', text_color = 'black', width = 171, height = 68)
        self.cap_value.place(x = 554, y = 476)

        self.upper_limit_label = CTkLabel(win, bg_color = '#CBD3F9', text="Cận trên", font = ('Times new roman', 25), fg_color = "#D465F0", text_color = "black", corner_radius = 10, 
                                          width = 118, height = 68)
        self.upper_limit_label.place(x=410, y = 386)

        self.upper_limit_entry = CTkTextbox(win, bg_color = '#CBD3F9', font = ('Times new roman', 25), fg_color = "#F8F3F9", text_color = "black", corner_radius = 10, width = 171, height = 68)
        self.upper_limit_entry.place(x=554, y = 386)

        self.lower_limit_label = CTkLabel(win, bg_color = '#CBD3F9', text="Cận dưới", font = ('Times new roman', 25), fg_color = "#D465F0", text_color = "black", corner_radius = 10, 
                                          width = 118, height = 68)
        self.lower_limit_label.place(x=751, y = 386)

        self.lower_limit_entry = CTkTextbox(win, bg_color = '#CBD3F9',  font = ('Times new roman', 25), fg_color = "#F8F3F9", text_color = "black", corner_radius = 10, width = 171, height = 68)
        self.lower_limit_entry.place(x=895, y = 386)


        integral_button = CTkButton(win, bg_color = '#CBD3F9', text = " Tích \n phân",command = self.compute_integral, font = ('Times new roman', 25), fg_color = "#D465F0", border_color = '#ccfff9', border_width = 2, text_color = "black", 
                                    corner_radius = 10, width = 165, height = 91, )
        integral_button.place(x = 341, y = 582)

        integral_button = CTkButton(win, bg_color = '#CBD3F9', text = " Đạo \n hàm",command = self.compute_derivative, font = ('Times new roman', 25), fg_color = "#D465F0", border_color = '#ccfff9', border_width = 2, text_color = "black", 
                                    corner_radius = 10, width = 165, height = 91, )
        integral_button.place(x = 341, y = 722)

        save_button = CTkButton(win, text = " Lưu kết quả ",command = self.storeResult, font = ('Times new roman', 25), fg_color = "#C6E8B1", border_color = '#12151c', border_width = 2, bg_color = '#B5EBF7', text_color = "black", corner_radius = 10, 
                                width = 165, height = 91, )
        save_button.place(x = 36, y = 245)

        help_button = CTkButton(win, text = " Hướng dẫn ",command = self.help, font = ('Times new roman', 25), fg_color = "#C6E8B1", border_color = '#12151c', border_width = 2, bg_color = '#B5EBF7', text_color = "black", corner_radius = 10, 
                                width = 165, height = 91, )
        help_button.place(x = 36, y = 141)

        self.outputIntegral = CTkTextbox(win, bg_color = '#CBD3F9', font = ('Times New Roman', 25), fg_color = "#F8F3F9", text_color = "black", width = 279, height = 231, corner_radius = 10)
        self.outputIntegral.place(x = 554, y = 582)

        self.outputDerivative = CTkTextbox(win, bg_color = '#CBD3F9', font = ('Times New Roman', 25), fg_color = "#F8F3F9", text_color = "black", width = 279, height = 231, corner_radius = 10)
        self.outputDerivative.place(x = 853, y = 582)

        self.iconUpload = CTkImage(light_image = Image.open('asset/iconUpload.jpg'), size = (83, 77))
        self.upload = CTkButton(win, image = self.iconUpload, corner_radius = 3, text = "", command = self.upFile, font = ('Times new roman', 25), fg_color = "#54AFC3", 
                                 width = 82, height = 76, hover = True, hover_color='#37a4de', border_color = '#37a4de', border_width = 3)
        self.upload.place(x = 1422, y = 18)

        self.file = CTkTextbox(win, bg_color = '#CBD3F9',  corner_radius = 10, font = ('Times New Roman', 25), height = 82, width = 186)
        self.file.configure(text_color = 'black', fg_color = "white", bg_color = '#B5EBF7')
        self.file.place(x = 1338, y = 152)

        if os.path.isfile("calIntegral/expression1.txt") == True and \
            os.path.isfile("calIntegral/resultIntegral.txt") == True and \
            os.path.isfile("calIntegral/resultDerivative.txt") == True:            
                pre_express = open("calIntegral/expression1.txt", "r")
                self.expression_entry.insert(INSERT, pre_express.read())

                pre_result1 = open("calIntegral/resultIntegral.txt", "r")
                text1="Tích phân của hàm số: \n" + pre_result1.read()
                self.outputIntegral.insert(INSERT, text1)

                pre_result2 = open("calIntegral/resultDerivative.txt", "r")
                text2 = "Đạo hàm của hàm số: \n" + pre_result2.read()
                self.outputDerivative.insert(INSERT, text2)
        if os.path.isfile("calIntegral/upperEdge.txt") == True and os.path.isfile("calIntegral/lowerEdge.txt") == True:
            upperEdg = open("calIntegral/upperEdge.txt", 'r')
            lowerEdg = open("calIntegral/lowerEdge.txt", 'r')
            self.upper_limit_entry.insert(INSERT, upperEdg.read())
            self.lower_limit_entry.insert(INSERT, lowerEdg.read())

    def selection(self,event):
        self.upper_limit_entry.delete('1.0', 'end')
        self.lower_limit_entry.delete("1.0", "end")
        self.upper_limit_entry.configure(state = "disabled" if self.radio_var1.get() == "indefinite" else "normal")
        self.upper_limit_entry.configure(fg_color = "#c5c9c9" if self.radio_var1.get() == "indefinite" else "#F8F3F9")
        self.lower_limit_entry.configure(state = "disabled" if self.radio_var1.get() == "indefinite" else "normal")
        self.lower_limit_entry.configure(fg_color = "#c5c9c9" if self.radio_var1.get() == "indefinite" else "#F8F3F9")

    def upFile(self):
        file_path = filedialog.askopenfilename()
        if file_path:
            file_name = file_path.split("/")[-1]
            if (file_name.split(".")[-1] != "txt"):
                messagebox.showwarning('Định dạng file không đúng', 'File được upload phải có định dạng .txt')
            else:
                try:
                    express = open(file_path, "r").read()
                    if express.strip() == "":
                        messagebox.showwarning('File có vấn đề!', 'File rỗng!')
                    self.file.delete("1.0", "end")
                    self.file.insert(INSERT, file_name)
                    self.expression_entry.delete("1.0", "end")
                    self.expression_entry.insert(INSERT, express)
                except:
                    if UnicodeDecodeError:
                        messagebox.showwarning("Lỗi File!", "Không đọc được nội dung trong file!")

    def compute_integral(self):
        try:
            expr = open("calIntegral/expression1.txt", "w")
            result1 = open("calIntegral/resultIntegral.txt", "w")
            self.outputIntegral.delete("1.0", "end")
            expression = self.expression_entry.get("1.0", "end")
            x = sym.symbols('x')
            expr.write(expression)
            if self.radio_var1.get() != "indefinite" and self.radio_var1.get() != "definite":
                messagebox.showwarning("Warning!", "Chọn loại tích phân trước khi tính!")
            else:
                if self.radio_var1.get() == "indefinite":
                    integral = sym.integrate(expression, x)
                    text = "\n  Tích phân của hàm số:\n    " + str(integral)
                    self.outputIntegral.insert(INSERT, text)
                    result1.write(str(integral))
                if self.radio_var1.get() == "definite":
                    upperEdge = open('calIntegral/upperEdge.txt', 'w')
                    lowerEdge = open('calIntegral/lowerEdge.txt', 'w')
                    a = self.lower_limit_entry.get("1.0", "end")
                    b = self.upper_limit_entry.get("1.0", "end")
                    if (a != "-oo\n" and a != "oo\n") and (b != "oo\n" and b != "-oo\n"):
                        a = str(a).replace("\n", "")
                        b = str(b).replace("\n", "")
                        a = str(a).replace("pi", "math.pi")
                        a = str(a).replace("e", "math.e")
                        b = str(b).replace("pi", "math.pi")
                        b = str(b).replace("e", "math.e")
                        integral = sym.integrate(expression, (x, float(eval(a)), float(eval(b))))
                        text = "\n  Tích phân của hàm số:\n    " + str(integral)
                    else:
                        list3 = ['-oo\n', '+oo\n', 'oo\n']
                        list4 = [-math.inf, math.inf, math.inf]
                        for i in range(len(list3)):
                            if list3[i] == a:
                                a = list4[i]
                            if list3[i] == b:
                                b = list4[i]
                        integral = sym.integrate(expression, (x, a, b))
                        text = "\n  Tích phân của hàm số:\n    " + str(integral)
                    self.outputIntegral.insert(INSERT, text)
                    upperEdge.write(str(b))
                    lowerEdge.write(str(a))
                    result1.write(str(integral))
                
        except:
            if SyntaxError or TypeError or sym.core.sympify.SympifyError:
                messagebox.showerror('Error', 'Kiểm tra lại đầu vào!')
                
    # xem ky phan nay  
    def compute_derivative(self):
        try:
            result2 = open("calIntegral/resultDerivative.txt", "w")
            self.outputDerivative.delete("1.0", "end")
            expression = self.expression_entry.get("1.0", "end").strip()
            x = sym.symbols('x')
            if self.radio_var2.get() == 'cấp 1':
                derivative = sym.diff(expression, x)
            if self.radio_var2.get() == 'cấp 2':
                derivative1 = sym.diff(expression, x)
                derivative = sym.diff(derivative1, x)
            if self.radio_var2.get() == 'cấp 3':
                derivative2 = sym.diff(expression, x)
                derivative1 = sym.diff(derivative2, x)
                derivative = sym.diff(derivative1, x)
            text="\n  Đạo hàm của hàm số:\n     " + str(derivative)
            self.outputDerivative.insert(INSERT, text)
            result2.write(str(derivative)) 
        except:
            messagebox.showerror('Error', 'Kiểm tra lại đầu vào!')

    def storeResult(self):
        try:
            doc = Document()
            text1 = f'-- Hàm số: {self.expression_entry.get("1.0", "end")}'
            doc.add_paragraph(text1)
            if self.radio_var1.get() == "indefinite":
                text2 = f'-- Tích phân không xác định của hàm số là: {self.outputIntegral.get("1.0", "end")}'
                doc.add_paragraph(text2)
            else:
                text3 = f'-- Tích phân xác định của hàm số là: {self.outputIntegral.get("1.0", "end")}\n'
                doc.add_paragraph(text3) 
                text4 = f'-- Với cận trên b = {self.upper_limit_entry.get("1.0", "end")}, cận dưới a = {self.lower_limit_entry.get("1.0", "end")}'
                doc.add_paragraph(text4)
            text5 = f'-- {self.outputDerivative.get("1.0", "end")}'
            doc.add_paragraph(text5)
            file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")], title="Save As")
            try:
                doc.save(f'{file_path}')
                messagebox.showinfo('Lưu file thành công!.',f'- Địa chỉ file: {file_path}')
            except PermissionError:
                messagebox.showinfo('Lưu file không thành công!', 'File đang được mở! Hãy tắt file đi!')
        except:
            messagebox.showerror('Lưu file thất bại!', 'Xem lại dữ liệu đầu vào!')
    def help(self):
        file = "Help/HelpIntegral/helpIntegral.py"
        subprocess.call(['python', file])

if __name__ == "__main__":
    win = CTk()
    app = calIntegral(win)
    win.mainloop()

